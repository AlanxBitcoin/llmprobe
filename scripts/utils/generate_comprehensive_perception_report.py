from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "outputs"
REPORT = OUT / "四感知类别_原始与打乱顺序_第八层深度解读综合报告.docx"

CATEGORIES = [
    {"key": "color_words", "cn": "颜色", "en": "Color", "accent": "1F4E79"},
    {"key": "taste_words", "cn": "味道", "en": "Taste", "accent": "375623"},
    {"key": "sound_words", "cn": "声音", "en": "Sound", "accent": "7030A0"},
    {"key": "shape_words", "cn": "形状", "en": "Shape", "accent": "C55A11"},
]

ACCENT = "1F4E79"
ACCENT_ORANGE = "C55A11"
SOFT_BLUE = "EAF2F8"
SOFT_ORANGE = "FCE4D6"
SOFT_GREEN = "E2F0D9"
SOFT_PURPLE = "EDE7F6"
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x55, 0x65, 0x73)


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def count_words(path: Path) -> int:
    return sum(1 for line in path.read_text(encoding="utf-8").splitlines() if line.strip())


def rows_for(cat: str, shuffled: bool, filename: str) -> list[dict[str, str]]:
    suffix = "_shuffled" if shuffled else ""
    return load_csv(OUT / f"{cat}{suffix}" / filename)


def top_dim(cat: str, shuffled: bool, mode_file: str, group: str, dim: str) -> dict[str, str] | None:
    for row in rows_for(cat, shuffled, mode_file):
        if row.get("group") == group and row.get("dim") == dim:
            return row
    return None


def top_dims(cat: str, shuffled: bool, mode_file: str, group: str, n: int = 5) -> list[dict[str, str]]:
    return [row for row in rows_for(cat, shuffled, mode_file) if row.get("group") == group][:n]


def fmt(value: str | float | int, ndigits: int = 3) -> str:
    try:
        return f"{float(value):.{ndigits}f}"
    except (TypeError, ValueError):
        return str(value)


def safe_count(row: dict[str, str] | None) -> str:
    return row.get("appearance_count", "0") if row else "0"


def safe_mean(row: dict[str, str] | None) -> str:
    return fmt(row.get("mean_abs_value", 0.0) if row else 0.0)


def set_font(run, size=10.5, bold=False, color=TEXT) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_bg(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7DEE8", size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def para(doc: Document, text: str = "", size=10.5, bold=False, color=TEXT, before=0, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.23
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc: Document, text: str, level: int, color_hex: str = ACCENT) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(15 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 4)
    r = p.add_run(text)
    size = 16 if level == 1 else 12.3
    set_font(r, size=size, bold=True, color=RGBColor.from_string(color_hex))


def callout(doc: Document, title: str, body: str, fill=SOFT_BLUE, accent=ACCENT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, fill)
    set_cell_border(cell, accent, "10")
    set_cell_margins(cell, 145, 170, 145, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.8, bold=True, color=RGBColor.from_string(accent))
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=10.0, color=TEXT)
    para(doc, after=2)


def table(doc: Document, headers: list[str], data: list[list[str]], widths: list[float] | None = None, accent=ACCENT):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, accent)
        set_cell_border(cell, "FFFFFF")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=8.6, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ridx, row in enumerate(data):
        cells = tbl.add_row().cells
        for cidx, value in enumerate(row):
            cell = cells[cidx]
            set_cell_bg(cell, "FFFFFF" if ridx % 2 == 0 else "F8FAFC")
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cidx in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=8.2, color=TEXT)
    if widths:
        for row in tbl.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    para(doc, after=4)
    return tbl


def add_img(doc: Document, path: Path, caption: str, width=6.25) -> None:
    if not path.exists():
        para(doc, f"[图片未找到：{path}]", size=9.2, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(7)
    r = cp.add_run(caption)
    set_font(r, size=8.6, color=MUTED)


def footer(section) -> None:
    p = section.footer.paragraphs[0]
    p._element.clear_content()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("四感知类别原始与打乱顺序第八层深度解读综合报告")
    set_font(r, size=8.2, color=MUTED)


def setup_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.75)
        section.bottom_margin = Cm(1.65)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)
        footer(section)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    return doc


def add_new_page(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    footer(section)


def original_summary_rows() -> list[list[str]]:
    rows = []
    for cat in CATEGORIES:
        key = cat["key"]
        word_count = count_words(ROOT / "data" / f"{key}.txt")
        d4055_pw = top_dim(key, False, "per_word_dim_average_abs.csv", "max", "4055")
        d4055_pos = top_dim(key, False, "positional_dim_average_abs.csv", "max", "4055")
        d1800_pw = top_dim(key, False, "per_word_dim_average_abs.csv", "min", "1800")
        rows.append(
            [
                cat["cn"],
                str(word_count),
                f"{safe_count(d4055_pw)}/{word_count}, {safe_mean(d4055_pw)}",
                f"{safe_count(d4055_pos)}/{word_count}, {safe_mean(d4055_pos)}",
                f"{safe_count(d1800_pw)}/{word_count}, {safe_mean(d1800_pw)}",
                "核心感知轴稳定" if int(safe_count(d4055_pos)) == word_count else "基本稳定",
            ]
        )
    return rows


def order_effect_rows() -> list[list[str]]:
    rows = []
    for cat in CATEGORIES:
        key = cat["key"]
        word_count = count_words(ROOT / "data" / f"{key}.txt")
        o4055 = top_dim(key, False, "positional_dim_average_abs.csv", "max", "4055")
        s4055 = top_dim(key, True, "positional_dim_average_abs.csv", "max", "4055")
        o2261 = top_dim(key, False, "positional_dim_average_abs.csv", "max", "2261")
        s2261 = top_dim(key, True, "positional_dim_average_abs.csv", "max", "2261")
        o1162 = top_dim(key, False, "positional_dim_average_abs.csv", "min", "1162")
        s1162 = top_dim(key, True, "positional_dim_average_abs.csv", "min", "1162")
        rows.append(
            [
                cat["cn"],
                f"{safe_mean(o4055)} → {safe_mean(s4055)}",
                f"{safe_count(o2261)}/{word_count} → {safe_count(s2261)}/{word_count}",
                f"{safe_count(o1162)}/{word_count} → {safe_count(s1162)}/{word_count}",
                "方向稳定，强度小幅变化",
            ]
        )
    return rows


def lexical_suppression_rows(shuffled: bool = False) -> list[list[str]]:
    rows = []
    for cat in CATEGORIES:
        key = cat["key"]
        word_count = count_words(ROOT / "data" / f"{key}.txt")
        items = []
        for dim in ("2485", "1815"):
            pw = top_dim(key, shuffled, "per_word_dim_average_abs.csv", "max", dim)
            pos = top_dim(key, shuffled, "positional_dim_average_abs.csv", "max", dim)
            items.append(f"{dim}: {safe_count(pw)}→{safe_count(pos)}/{word_count}")
        rows.append([cat["cn"], "；".join(items), "序列上下文压制词法/形态噪声"])
    return rows


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("四感知类别原始与打乱顺序\n第八层深度解读综合报告")
    set_font(r, size=23, bold=True, color=RGBColor.from_string(ACCENT))
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(24)
    r = sp.add_run("颜色 / 味道 / 声音 / 形状  |  三种输入模式  |  Meta-Llama-3-8B-Instruct")
    set_font(r, size=12.3, color=RGBColor.from_string(ACCENT_ORANGE))
    callout(
        doc,
        "报告目标",
        "整合原始顺序四词库与打乱顺序四词库的第 8 层探针结果，解释哪些维度代表稳定感知语义，哪些维度来自词法表面特征，哪些维度来自列表上下文结构。",
        fill=SOFT_BLUE,
    )
    table(
        doc,
        ["实验组", "词库", "模式", "主要输出"],
        [
            ["原始顺序", "color/taste/sound/shape", "逐词、全量末 token、序列位置", "维度极值、平均绝对值、图表、视频"],
            ["打乱顺序", "shuffled_color/taste/sound/shape", "同上", "对比词序对上下文编码的影响"],
        ],
        widths=[2.5, 4.2, 4.2, 4.8],
    )


def add_executive_summary(doc: Document) -> None:
    add_new_page(doc)
    heading(doc, "1. 执行摘要", 1)
    callout(
        doc,
        "总判断",
        "第 8 层已经形成稳定的感知语义编码轴，但不是纯语义层；它同时混合词法表面特征和上下文列表结构。改变词序不会破坏核心感知轴，但会改变位置模式和全词末 token 模式中的上下文强度。",
        fill=SOFT_ORANGE,
        accent=ACCENT_ORANGE,
    )
    table(
        doc,
        ["发现", "证据", "解释"],
        [
            ["核心感知轴稳定", "dim4055 在四类原始位置模式均 100% 覆盖", "模型在第 8 层已能识别可感知属性词"],
            ["双极轴存在", "dim1800 在四类逐词模式均 100% 覆盖", "正负方向共同定位感知语义强度"],
            ["词法噪声被上下文压制", "dim2485/dim1815 从逐词高覆盖转为位置低覆盖", "序列语境稀释孤立 token 的词形信号"],
            ["列表结构维度涌现", "dim2261/2265/1162 在位置模式大面积出现", "模型识别到同类词枚举列表结构"],
            ["词序影响有限但存在", "逐词模式完全一致，位置模式小幅变化，全词末 token 更敏感", "语义轴稳健，局部上下文读数受顺序影响"],
        ],
        widths=[3.0, 5.0, 6.2],
    )


def add_original_analysis(doc: Document) -> None:
    heading(doc, "2. 原始顺序四词库：第 8 层感知语义结构", 1)
    para(
        doc,
        "原始顺序实验提供基准结果。逐词模式体现词本身的孤立语义；序列位置模式体现同类词列表上下文中的隐藏状态；全词末 token 模式只读最后一个 token，不能代表全部词的平均。",
    )
    table(
        doc,
        ["类别", "词数", "dim4055 逐词", "dim4055 位置", "dim1800 逐词", "判断"],
        original_summary_rows(),
        widths=[1.4, 1.0, 3.0, 3.0, 3.0, 2.8],
    )
    callout(
        doc,
        "核心解释",
        "dim4055 不宜被命名为“颜色维度”。它在颜色、味道、声音、形状四类中都稳定强激活，更像“具体感知属性词”的通用识别轴。dim1800 则与其构成负向对照轴。",
        fill=SOFT_BLUE,
    )
    table(
        doc,
        ["类别", "逐词 top max 维度", "逐词 top min 维度"],
        [
            [
                cat["cn"],
                "；".join(f"{r['dim']}({safe_count(r)},{safe_mean(r)})" for r in top_dims(cat["key"], False, "per_word_dim_average_abs.csv", "max", 4)),
                "；".join(f"{r['dim']}({safe_count(r)},{safe_mean(r)})" for r in top_dims(cat["key"], False, "per_word_dim_average_abs.csv", "min", 4)),
            ]
            for cat in CATEGORIES
        ],
        widths=[1.4, 6.5, 6.5],
    )


def add_three_layer_model(doc: Document) -> None:
    heading(doc, "3. 统一三层编码模型", 1)
    table(
        doc,
        ["层次", "代表维度", "判定依据", "功能解释"],
        [
            ["核心感知语义轴", "4055 / 290 / 1800 / 709", "跨类别、跨顺序、跨模式稳定", "识别“这是感知属性词”并定位强弱"],
            ["词法表面层", "2485 / 1815 / 912 / 1856 / 1731", "逐词模式强，位置模式被压制", "英文词形、词频、子词分割、孤立 token 噪声"],
            ["列表上下文结构层", "2261 / 2265 / 1162 / 2977 / 2116", "逐词近沉默，序列位置中涌现", "编码同类感知词被放在枚举列表中的结构信息"],
        ],
        widths=[3.0, 4.2, 4.0, 3.8],
    )
    para(doc, "词法表面维度在原始顺序中的抑制：", bold=True)
    table(doc, ["类别", "dim2485 / dim1815 逐词→位置", "解释"], lexical_suppression_rows(False), widths=[1.5, 6.0, 6.8])
    para(doc, "打乱顺序后，词法抑制仍然存在：", bold=True)
    table(doc, ["类别", "dim2485 / dim1815 逐词→位置", "解释"], lexical_suppression_rows(True), widths=[1.5, 6.0, 6.8])


def add_order_effect_analysis(doc: Document) -> None:
    add_new_page(doc)
    heading(doc, "4. 改变单词顺序后的影响", 1)
    callout(
        doc,
        "顺序影响结论",
        "改变词库顺序不会影响逐词模式，因为每个词仍单独输入；但会影响序列位置模式和全词末 token 模式，因为每个词的左侧上下文发生变化。实验显示：核心感知轴保持稳定，结构维度仍会涌现，但覆盖率和均值会小幅重排。",
        fill=SOFT_GREEN,
        accent="548235",
    )
    table(
        doc,
        ["类别", "dim4055 位置均值 原始→打乱", "dim2261 覆盖 原始→打乱", "dim1162 覆盖 原始→打乱", "解释"],
        order_effect_rows(),
        widths=[1.4, 3.0, 3.1, 3.1, 4.2],
    )
    table(
        doc,
        ["模式", "顺序敏感性", "原因", "建议解释方式"],
        [
            ["逐词模式", "不敏感", "每个单词单独输入，没有列表上下文", "可作为词本身语义基准"],
            ["全词末 token", "最敏感", "只读最后 token，左侧上下文和末词身份都很关键", "不能当作全词库平均表征"],
            ["序列位置模式", "中等敏感", "每个词的左侧上下文改变，但列表语义场仍存在", "最适合研究上下文如何重塑词表示"],
        ],
        widths=[2.3, 2.0, 5.0, 5.2],
        accent=ACCENT_ORANGE,
    )


def add_category_sections(doc: Document) -> None:
    add_new_page(doc)
    heading(doc, "5. 分类别深度解读", 1)
    for cat in CATEGORIES:
        key = cat["key"]
        cn = cat["cn"]
        accent = cat["accent"]
        heading(doc, f"5.{CATEGORIES.index(cat)+1} {cn}词：原始与打乱顺序对照", 2, accent)
        pw4055 = top_dim(key, False, "per_word_dim_average_abs.csv", "max", "4055")
        pos4055 = top_dim(key, False, "positional_dim_average_abs.csv", "max", "4055")
        spos4055 = top_dim(key, True, "positional_dim_average_abs.csv", "max", "4055")
        para(
            doc,
            f"{cn}词的 dim4055 逐词均值为 {safe_mean(pw4055)}，原始位置均值为 {safe_mean(pos4055)}，打乱后位置均值为 {safe_mean(spos4055)}。"
            "这说明核心感知语义轴在顺序变化后仍然存在，但上下文会改变其强度。",
        )
        table(
            doc,
            ["实验", "位置模式 top max", "位置模式 top min"],
            [
                [
                    "原始顺序",
                    "；".join(f"{r['dim']}({safe_count(r)},{safe_mean(r)})" for r in top_dims(key, False, "positional_dim_average_abs.csv", "max", 5)),
                    "；".join(f"{r['dim']}({safe_count(r)},{safe_mean(r)})" for r in top_dims(key, False, "positional_dim_average_abs.csv", "min", 5)),
                ],
                [
                    "打乱顺序",
                    "；".join(f"{r['dim']}({safe_count(r)},{safe_mean(r)})" for r in top_dims(key, True, "positional_dim_average_abs.csv", "max", 5)),
                    "；".join(f"{r['dim']}({safe_count(r)},{safe_mean(r)})" for r in top_dims(key, True, "positional_dim_average_abs.csv", "min", 5)),
                ],
            ],
            widths=[2.0, 6.0, 6.0],
            accent=accent,
        )
        add_img(doc, OUT / key / "three_mode_comparison_chart.png", f"{cn}词原始顺序三模式维度对比", width=6.1)
        add_img(doc, OUT / f"{key}_shuffled" / "three_mode_comparison_chart.png", f"{cn}词打乱顺序三模式维度对比", width=6.1)


def add_methodology_and_conclusions(doc: Document) -> None:
    add_new_page(doc)
    heading(doc, "6. 方法论反思与最终结论", 1)
    callout(
        doc,
        "最重要的方法论提醒",
        "单个维度不能直接等同于一个人类概念。更稳妥的解释是：某维度在某类词、某种模式、某类上下文中稳定高激活，因此近似承担某种功能方向。",
        fill=SOFT_PURPLE,
        accent="7030A0",
    )
    table(
        doc,
        ["结论编号", "结论", "研究含义"],
        [
            ["C1", "感知语义在第 8 层已高度结晶", "四类词均有稳定核心轴，说明模型中层已形成感知词识别机制"],
            ["C2", "第 8 层仍不是纯语义层", "词法表面维度仍显著存在，需要与语义维度区分"],
            ["C3", "列表上下文会产生独立结构编码", "模型不只理解单词，也识别“同类词列表”这种语用结构"],
            ["C4", "词序改变不会推翻核心结论", "核心轴稳健，顺序主要影响上下文强度与末 token 读数"],
            ["C5", "模式 1 与模式 3 的对照最有解释力", "可用于分离词本身语义与列表上下文效应"],
            ["C6", "未来应做跨层验证", "低层可能词法更强，高层可能语义更纯，需第 1/4/16/24/32 层对照"],
        ],
        widths=[1.6, 5.2, 7.2],
    )
    para(
        doc,
        "下一步建议：围绕 dim4055/1800/290/2485/1815/2261/2265/1162 建立“维度功能分层表”，再扩展到更多语义类别和更多层数。"
        "如果某个维度在不同类别、不同顺序、不同模式中表现一致，才更接近稳定机制；如果只在某个列表或末 token 中出现，则更可能是上下文或位置效应。",
    )


def build() -> None:
    doc = setup_doc()
    add_cover(doc)
    add_executive_summary(doc)
    add_original_analysis(doc)
    add_three_layer_model(doc)
    add_order_effect_analysis(doc)
    add_category_sections(doc)
    add_methodology_and_conclusions(doc)
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    build()
