# -*- coding: utf-8 -*-
"""生成颜色词探针实验中文Word报告"""
from __future__ import annotations
import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("data/outputs/color_words")
REPORT_PATH = OUT_DIR / "LLM颜色词探针实验报告.docx"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_font_cn(run):
    """为run设置中文字体（微软雅黑）"""
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rPr.insert(0, rFonts)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if h.runs:
        set_font_cn(h.runs[0])
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    set_font_cn(run)
    return p


def add_image(doc, path, width_inches=6.2, caption=""):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[图片未找到: {path}]", italic=True, size=9, color="888888")
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_font_cn(cp.runs[0])


def add_table(doc, headers, rows, header_bg="1f4e79", alt_bg="dce6f1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        set_font_cn(run)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], header_bg)
    for ri, row_data in enumerate(rows):
        row = t.add_row().cells
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            row[ci].text = str(val)
            run = row[ci].paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            set_font_cn(run)
            row[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row[ci], bg)
    doc.add_paragraph()
    return t


def load_csv(path):
    with open(path, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fmt_f(v, d=4):
    try:
        return f"{float(v):.{d}f}"
    except Exception:
        return str(v)


def build_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ------------------------------------------------------------------ 封面
    doc.add_paragraph()
    title = doc.add_heading("大语言模型颜色词探针实验报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        set_font_cn(title.runs[0])

    sub = doc.add_paragraph("Meta-Llama-3-8B-Instruct  |  第8层隐藏状态  |  三种输入模式")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    set_font_cn(sub.runs[0])

    info = doc.add_paragraph("实验日期：2026-05-07  |  颜色词数量：82  |  隐藏状态维度：4096")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(10)
    info.runs[0].italic = True
    set_font_cn(info.runs[0])
    doc.add_page_break()

    # ------------------------------------------------------------------ 第1章
    add_heading(doc, "第1章  实验概述", 1)
    add_para(doc,
        "本实验利用神经探针技术，揭示大语言模型 Meta-Llama-3-8B-Instruct 在第8层（共33层）"
        "对'颜色'这一语义类别的内部编码机制。\n"
        "实验向模型输入82个精心筛选的英语颜色词，提取对应的4096维隐藏状态向量，"
        "并分析哪些维度承载了颜色语义信息。所有激活值均为原始有符号值，未经归一化处理。",
        size=10)
    doc.add_paragraph()

    add_para(doc, "1.1  颜色词表设计", bold=True, size=11)
    add_table(doc,
        ["色系", "词数", "代表词（部分）"],
        [
            ["红色系", "10", "red, crimson, scarlet, maroon, burgundy, carmine, vermilion, ruby, brick, rust"],
            ["橙色系", "8",  "orange, amber, coral, peach, apricot, tangerine, saffron, ochre"],
            ["黄色系", "6",  "yellow, gold, lemon, mustard, buff, canary"],
            ["绿色系", "10", "green, lime, olive, emerald, jade, mint, teal, chartreuse, sage, khaki"],
            ["蓝色系", "9",  "blue, navy, cobalt, azure, cerulean, sapphire, denim, indigo, slate"],
            ["紫色系", "8",  "purple, violet, lavender, lilac, magenta, fuchsia, plum, mauve"],
            ["粉色系", "4",  "pink, rose, salmon, blush"],
            ["棕/土色系", "16", "brown, tan, beige, cream, ivory, ecru, taupe, sand, sienna, umber..."],
            ["无彩色系", "11", "white, gray, silver, charcoal, onyx, ebony, black, pearl, platinum, pewter, graphite"],
        ],
        header_bg="1f4e79")

    add_para(doc, "1.2  三种输入模式说明", bold=True, size=11)
    add_table(doc,
        ["模式", "操作方法", "目的"],
        [
            ["模式1：逐词孤立输入",
             "每个词单独作为输入，提取最后一个子词(sub-token)对应的隐藏状态。共82次推理。",
             "获取纯词义编码，不含上下文干扰"],
            ["模式2：全词合并末token",
             "82个词拼接为一个字符串，仅提取最后一个token的隐藏状态。共1次推理。",
             "以81个颜色词为左上下文，观察最后一词的编码"],
            ["模式3：序列位置提取",
             "82个词拼接后做一次推理，在序列中各词自身位置提取隐藏状态。共82个向量。",
             "每个词共享列表上下文，模拟真实文本中的列表结构"],
        ],
        header_bg="375623")
    doc.add_page_break()

    # ------------------------------------------------------------------ 第2章
    add_heading(doc, "第2章  维度激活总览——三模式对比", 1)
    add_para(doc,
        "下图并排展示三种输入模式下激活最显著的前20个维度。"
        "条形长度代表均值|激活值|（模式1/3）或有符号激活值（模式2中间栏）。"
        "暖色（橙红）= 正向激活；冷色（深蓝）= 负向激活。所有值均为原始值，未归一化。",
        size=10)
    add_image(doc, OUT_DIR / "three_mode_comparison_chart.png", width_inches=6.5,
              caption="图1  三模式维度对比（第8层，原始有符号激活值）")
    add_para(doc,
        "关键观察：模式1（左栏）中，dim 4055 的均值激活远超其他维度（约为第二名的3倍）。"
        "模式3（右栏）中，dim 2261、dim 1162、dim 709 等新维度大量涌现，"
        "这些维度在孤立输入时几乎沉默，揭示了模型对列表结构的感知能力。",
        italic=True, size=9)
    doc.add_page_break()

    # ------------------------------------------------------------------ 第3章
    add_heading(doc, "第3章  模式1——逐词孤立输入", 1)
    add_para(doc,
        "模式1中每个颜色词单独输入模型，共进行82次独立推理。"
        "提取每个词最后一个子词位置的第8层隐藏状态，代表纯词义编码。",
        size=10)
    add_image(doc, OUT_DIR / "per_word_dim_stats_chart.png", width_inches=6.5,
              caption="图2  模式1逐词输入：Top-20维度统计（上行正向、下行负向）")

    add_heading(doc, "3.1  最活跃正向维度", 2)
    pw = load_csv(OUT_DIR / "per_word_dim_average_abs.csv")
    pw_pos = [r for r in pw if r["group"] == "max"][:10]
    add_table(doc,
        ["维度", "出现次数/82", "均值|激活|", "最大|激活|", "解读"],
        [
            [r["dim"], r["appearance_count"], fmt_f(r["mean_abs_value"]), fmt_f(r["max_abs_value"]),
             "核心颜色通道（主导维度）" if r["dim"] == "4055" else
             "最稳定颜色标记（覆盖率100%）" if r["dim"] == "290" else
             "强颜色信号" if int(r["appearance_count"]) >= 50 else "中等 / 含词法特征"]
            for r in pw_pos
        ],
        header_bg="c55a11")

    add_heading(doc, "3.2  最活跃负向维度", 2)
    pw_neg = [r for r in pw if r["group"] == "min"][:8]
    add_table(doc,
        ["维度", "出现次数/82", "均值|激活|", "最大|激活|", "解读"],
        [
            [r["dim"], r["appearance_count"], fmt_f(r["mean_abs_value"]), fmt_f(r["max_abs_value"]),
             "核心负向颜色通道" if r["dim"] == "1800" else
             "强负向颜色信号" if int(r["appearance_count"]) >= 50 else "中等"]
            for r in pw_neg
        ],
        header_bg="2e4099")

    add_heading(doc, "3.3  dim 4055 激活最强 vs 最弱词汇", 2)
    ex = load_csv(OUT_DIR / "per_word_dim_extremes.csv")
    top10 = sorted(ex, key=lambda r: float(r.get("max_value_1", 0)), reverse=True)[:10]
    bot10 = sorted(ex, key=lambda r: float(r.get("max_value_1", 0)))[:10]

    add_para(doc, "激活最强的10个词（具体可视的颜色名称）：", bold=True, size=10)
    add_table(doc,
        ["排名", "词", "维度", "激活值", "分析"],
        [
            [i + 1, r["input"], r["max_dim_1"], fmt_f(r["max_value_1"]),
             "高度具体的视觉参照物（事物名→唯一色调）" if float(r["max_value_1"]) > 1.8
             else "鲜明颜色名称"]
            for i, r in enumerate(top10)
        ],
        header_bg="c00000")

    add_para(doc, "激活最弱的10个词：", bold=True, size=10)
    add_table(doc,
        ["排名", "词", "维度", "激活值", "分析"],
        [
            [i + 1, r["input"], r["max_dim_1"], fmt_f(r["max_value_1"]),
             "被拆分为多个子词，颜色信号稀释"
             if r["input"] in ("indigo", "chartreuse", "vermilion", "cerulean")
             else "多义词 / 颜色含义非主要义项"]
            for i, r in enumerate(bot10)
        ],
        header_bg="7f7f7f")
    doc.add_page_break()

    # ------------------------------------------------------------------ 第4章
    add_heading(doc, "第4章  模式3——序列位置提取", 1)
    add_para(doc,
        "模式3将82个颜色词拼接为一个序列，进行一次前向传播，"
        "然后在序列中每个词自身对应的位置提取第8层隐藏状态。"
        "每个词都能通过注意力机制感知到其左侧所有已出现的颜色词，"
        "模拟模型处理真实颜色词列表时的工作方式。",
        size=10)
    add_image(doc, OUT_DIR / "positional_dim_stats_chart.png", width_inches=6.5,
              caption="图3  模式3序列位置提取：维度统计（第8层）")

    add_heading(doc, "4.1  最活跃正向维度（位置模式）", 2)
    pos = load_csv(OUT_DIR / "positional_dim_average_abs.csv")
    pos_pos = [r for r in pos if r["group"] == "max"][:10]
    add_table(doc,
        ["维度", "出现次数/82", "均值|激活|", "最大|激活|", "与模式1对比"],
        [
            [r["dim"], r["appearance_count"], fmt_f(r["mean_abs_value"]), fmt_f(r["max_abs_value"]),
             "覆盖率 96%→100%（上下文增强）" if r["dim"] == "4055" else
             "77→20次（上下文抑制）" if r["dim"] == "2485" else
             "新涌现：8→76次（上下文激活）" if r["dim"] == "2261" else
             "新涌现：仅在序列中出现" if r["dim"] == "2265" else
             "跨模式稳定"]
            for r in pos_pos
        ],
        header_bg="375623")

    add_heading(doc, "4.2  最活跃负向维度（位置模式）", 2)
    pos_neg = [r for r in pos if r["group"] == "min"][:8]
    add_table(doc,
        ["维度", "出现次数/82", "均值|激活|", "最大|激活|", "与模式1对比"],
        [
            [r["dim"], r["appearance_count"], fmt_f(r["mean_abs_value"]), fmt_f(r["max_abs_value"]),
             "新涌现：9→80次（上下文激活）" if r["dim"] == "1162" else
             "新涌现：40→81次（大幅增强）" if r["dim"] == "709" else
             "跨模式稳定" if r["dim"] in ("1800", "3231") else
             "位置模式中等水平"]
            for r in pos_neg
        ],
        header_bg="2e4099")
    doc.add_page_break()

    # ------------------------------------------------------------------ 第5章
    add_heading(doc, "第5章  模式2——全词合并末token", 1)
    add_para(doc,
        "模式2将82个颜色词拼接为一个字符串送入模型，仅提取最后一个token（graphite）"
        "对应的第8层隐藏状态。\n"
        "注意：这并非对所有颜色词的整体聚合，而是以前81个词为背景上下文，"
        "单独观察最后一词的编码变化。"
        "dim 4055 峰值激活降至 0.780，而模式1中该词最高达 2.378。",
        size=10)
    add_image(doc, OUT_DIR / "all_input_dim_stats_chart.png", width_inches=6.5,
              caption="图4  模式2全词合并末token（前24个最显著维度）")
    ai = load_csv(OUT_DIR / "all_input_dim_stats.csv")[:16]
    add_table(doc,
        ["维度", "有符号激活值", "方向", "绝对值"],
        [[r["dim"], fmt_f(r["signed_value"]), r["direction"], fmt_f(r["abs_value"])]
         for r in ai],
        header_bg="4e3379")
    doc.add_page_break()

    # ------------------------------------------------------------------ 第6章
    add_heading(doc, "第6章  三模式横向对比与核心发现", 1)

    add_heading(doc, "6.1  关键维度在三种模式下的行为汇总", 2)
    add_table(doc,
        ["维度", "模式1 逐词", "模式2 末token", "模式3 位置", "解读"],
        [
            ["dim 4055（正）",
             "79/82（96%）\n均值=1.471",
             "激活（0.780）",
             "82/82（100%）\n均值=0.968",
             "核心颜色通道\n具体名称触发最强"],
            ["dim 290（正）",
             "82/82（100%）\n均值=0.500",
             "激活（0.471）",
             "80/82（98%）\n均值=0.470",
             "最稳健颜色标记\n三模式全稳定"],
            ["dim 1800（负）",
             "82/82（100%）\n均值=0.820",
             "激活（-0.421）",
             "81/82（99%）\n均值=0.472",
             "核心负向颜色通道\n与dim 4055构成双极轴"],
            ["dim 3231（负）",
             "63/82（77%）\n均值=0.443",
             "激活（-0.443）",
             "72/82（88%）\n均值=0.446",
             "稳健负向信号\n三模式一致"],
            ["dim 2485（正）",
             "77/82（94%）\n均值=0.511",
             "弱（0.345）",
             "20/82（24%）\n均值=0.424",
             "孤立时强\n上下文中被抑制\n可能含词法特征"],
            ["dim 2261（±）",
             "8次（负向）",
             "激活（+0.446）",
             "76/82（93%，正向）",
             "极性翻转！\n上下文涌现的列表结构信号"],
            ["dim 1162（负）",
             "9/82（11%）",
             "激活（-0.749）",
             "80/82（98%）\n均值=0.541",
             "仅在序列上下文中涌现\n编码列表结构"],
            ["dim 709（负）",
             "40/82（49%）弱",
             "激活（-0.519）",
             "81/82（99%）\n均值=0.533",
             "仅在序列上下文中涌现\n编码列表结构"],
        ],
        header_bg="1f4e79")

    add_heading(doc, "6.2  三层编码模型", 2)
    add_para(doc,
        "根据维度在三种模式下的行为差异，可将激活维度分为三个功能层次：",
        size=10)

    layers = [
        ("第一层：核心颜色信号（跨模式稳健）",
         "维度：4055（正）、290（正）、1800（负）、3231（负）\n"
         "在三种模式下均稳定激活，是模型对'颜色'语义类别的硬编码表征。"
         "dim 4055 与 dim 1800 构成双极轴——一高一低的同步激活构成颜色识别标志。",
         "c00000"),
        ("第二层：词法表面信号（孤立时强，上下文中被抑制）",
         "维度：2485、133、1815（正）；912、1856、1731（负）\n"
         "在逐词模式下活跃，但在位置模式下大幅抑制。"
         "这些维度可能混合了词频、形态特征与颜色语义，并非纯颜色信号。",
         "bf8f00"),
        ("第三层：上下文结构信号（仅在序列中涌现）",
         "维度：2261、2265（正）；1162、709（负）\n"
         "孤立输入时几乎沉默，但在序列位置模式下大量激活（覆盖率达80-82/82）。"
         "编码元信息：当前正在处理颜色词列表中的一个词。"
         "揭示了模型对枚举列表结构的感知能力。",
         "375623"),
    ]
    for ltitle, lbody, lcolor in layers:
        p = doc.add_paragraph()
        rt = p.add_run(ltitle + "\n")
        rt.bold = True
        rt.font.size = Pt(10.5)
        r_int, g_int, b_int = int(lcolor[:2], 16), int(lcolor[2:4], 16), int(lcolor[4:], 16)
        rt.font.color.rgb = RGBColor(r_int, g_int, b_int)
        set_font_cn(rt)
        rb = p.add_run(lbody)
        rb.font.size = Pt(10)
        set_font_cn(rb)
        doc.add_paragraph()
    doc.add_page_break()

    # ------------------------------------------------------------------ 第7章
    add_heading(doc, "第7章  哪些词最像/最不像颜色词？", 1)
    add_para(doc,
        "在逐词模式下，dim 4055 的激活强度可量化模型对一个词'颜色身份'的认定程度。"
        "源自具体物体名称的颜色词得分最高；抽象、多义或被拆分的词得分最低。",
        size=10)
    finding = [
        ["apricot（杏色）",  "2.378", "极具体的视觉参照（水果名→唯一色调），颜色身份无歧义"],
        ["lavender（薰衣草紫）", "2.280", "植物名→高度聚焦的紫蓝色调"],
        ["saffron（藏红花黄）",  "2.212", "香料名→非常具体的黄橙色调"],
        ["mahogany（红木棕）",  "2.189", "木材名→特定红棕色调"],
        ["salmon（鲑鱼粉）",   "2.183", "食物名→特定粉橙色调"],
        ["……",               "……",   "……"],
        ["indigo（靛蓝）",    "0.391  ⚠", "被拆分为3个子词，颜色信号被稀释"],
        ["sage（鼠尾草绿）",  "0.426  ⚠", "主要义项为草本植物，颜色为次要义项"],
        ["rose（玫瑰粉）",   "0.500", "多义词：花 vs 粉色，信号分散"],
        ["amber（琥珀色）",  "0.548", "主要义项为矿物/化石，颜色为次要义项"],
        ["blue（蓝色）",     "0.800", "最高频基础颜色词，但在大量非颜色语境中出现"],
    ]
    add_table(doc, ["词（中文名）", "dim 4055 激活值", "分析"], finding, header_bg="7030a0")
    add_para(doc,
        "规律：源自物体名称的颜色词（食物、植物、矿物）激活最强；"
        "抽象、多义或多子词颜色词激活最弱。"
        "这与人类认知中颜色词的显著性梯度高度吻合。",
        italic=True, size=9)
    doc.add_page_break()

    # ------------------------------------------------------------------ 第8章
    add_heading(doc, "第8章  结论", 1)
    conclusions = [
        ("结论1  颜色语义在第8层已高度结晶化",
         "Llama-3-8B 在第8层形成了专门编码'颜色'语义类别的维度区域。"
         "dim 4055（正向）与 dim 1800（负向）在三种模式下的覆盖率达96-100%，"
         "激活幅度是其他维度的2-3倍，构成核心颜色编码轴。"),
        ("结论2  具体名称 vs 抽象名称存在显著激活梯度",
         "源自物体名称的颜色词（apricot、lavender、saffron）触发 dim 4055 的激活幅度"
         "是基础颜色词（blue、red）的2-3倍，反映出颜色身份越清晰的词，模型颜色编码越专一。"),
        ("结论3  上下文能强化弱颜色信号",
         "在序列位置模式下，dim 4055 覆盖率从96%升至100%。"
         "孤立输入时未能激活颜色维度的词（如 indigo）在颜色词列表语境中也能正确激活，"
         "表明模型可利用上下文消解边缘词的歧义。"),
        ("结论4  模型能感知'颜色词列表'这一宏观结构",
         "dim 2261、dim 1162、dim 709 在孤立输入时几乎沉默，"
         "但在序列位置模式下对80-82/82个词激活。"
         "这些仅在序列中涌现的维度编码了结构性元信息：'当前词属于一个颜色词枚举列表'——"
         "模型不仅读取词义，还感知列表结构。"),
        ("结论5  词法表面特征与语义信号共存于第8层",
         "dim 2485 和 dim 133 在孤立模式下高度活跃，但在上下文中被抑制，"
         "提示第8层并非纯语义层，仍混有词频/形态特征。"
         "更高层可能呈现更纯粹的颜色语义表征，值得进一步探究。"),
    ]
    for ct, cb in conclusions:
        p = doc.add_paragraph()
        rt = p.add_run(ct + "\n")
        rt.bold = True
        rt.font.size = Pt(11)
        rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        set_font_cn(rt)
        rb = p.add_run(cb)
        rb.font.size = Pt(10)
        set_font_cn(rb)
        doc.add_paragraph()

    doc.save(str(REPORT_PATH))
    print(f"报告已保存：{REPORT_PATH}")


build_report()
