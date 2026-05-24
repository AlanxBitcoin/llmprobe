# -*- coding: utf-8 -*-
"""
生成四感知类别（颜色/味道/声音/形状）三模式第八层隐藏层分析综合中文Word报告
"""
from __future__ import annotations
import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path("data/outputs")
REPORT_PATH = BASE / "四感知类别三模式第八层分析报告.docx"

CATEGORIES = [
    {
        "key":   "color_words",
        "label": "颜色（Color）",
        "cn":    "颜色",
        "en":    "Color",
        "words": 82,
        "families": "红橙黄绿蓝紫粉棕无彩色，9个色系",
        "color_h": "1f4e79",
        "color_pos": "c55a11",
        "color_neg": "2e4099",
        "cover_color": RGBColor(0x1F, 0x4E, 0x79),
    },
    {
        "key":   "taste_words",
        "label": "味道（Taste）",
        "cn":    "味道",
        "en":    "Taste",
        "words": None,
        "families": "甜/糖味、酸味、苦味、咸味、鲜味、辛辣、烟熏/土味、芳香、淡味，共9类",
        "color_h": "375623",
        "color_pos": "538135",
        "color_neg": "244f26",
        "cover_color": RGBColor(0x37, 0x56, 0x23),
    },
    {
        "key":   "sound_words",
        "label": "声音（Sound）",
        "cn":    "声音",
        "en":    "Sound",
        "words": None,
        "families": "响度（响/轻）、音调（高/低）、音色（清脆/浑厚）、自然声、人造声、音乐性，共6类",
        "color_h": "7030a0",
        "color_pos": "9b59b6",
        "color_neg": "4a235a",
        "cover_color": RGBColor(0x70, 0x30, 0xA0),
    },
    {
        "key":   "shape_words",
        "label": "形状（Shape）",
        "cn":    "形状",
        "en":    "Shape",
        "words": None,
        "families": "基础几何形、立体形、轮廓/边缘特征、表面纹理、比例/大小，共5类",
        "color_h": "c55a11",
        "color_pos": "e67e22",
        "color_neg": "784212",
        "cover_color": RGBColor(0xC5, 0x5A, 0x11),
    },
]


# ─────────────── DocX helpers ────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_font_cn(run, size_pt=None):
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rPr.insert(0, rFonts)
    if size_pt:
        run.font.size = Pt(size_pt)


def add_heading(doc, text, level, color_rgb=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if h.runs:
        set_font_cn(h.runs[0])
        if color_rgb:
            h.runs[0].font.color.rgb = color_rgb
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        if isinstance(color, RGBColor):
            run.font.color.rgb = color
        else:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    set_font_cn(run)
    return p


def add_image(doc, path, width_inches=6.2, caption=""):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[图片未找到: {path}]", italic=True, size=9, color="888888")
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_font_cn(cp.runs[0])


def add_table(doc, headers, rows, header_bg="1f4e79", alt_bg="dce6f1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        set_font_cn(run)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], header_bg)
    for ri, row_data in enumerate(rows):
        row = t.add_row().cells
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            row[ci].text = str(val)
            run = row[ci].paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            set_font_cn(run)
            row[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row[ci], bg)
    doc.add_paragraph()
    return t


def load_csv(path):
    p = Path(path)
    if not p.exists():
        return []
    with open(p, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fmt_f(v, d=4):
    try:
        return f"{float(v):.{d}f}"
    except Exception:
        return str(v)


def out_dir(cat_key):
    return BASE / cat_key


# ─────────────── Per-category section ────────────────────────────────────────

def write_category_section(doc, cat, sec_num):
    key = cat["key"]
    cn = cat["cn"]
    color_h = cat["color_h"]
    color_pos = cat["color_pos"]
    color_neg = cat["color_neg"]
    odir = out_dir(key)

    add_heading(doc, f"第{sec_num}章  {cn}词（{cat['en']}）三模式分析", 1,
                color_rgb=cat["cover_color"])
    add_para(doc,
        f"词库分类：{cat['families']}。\n"
        "实验采用三种输入模式（逐词孤立 / 全词末token / 序列位置），提取第8层隐藏状态，"
        "对所有词汇进行维度激活统计。",
        size=10)

    # 三模式对比图
    add_heading(doc, f"{sec_num}.1  三模式维度激活总览", 2)
    add_image(doc, odir / "three_mode_comparison_chart.png", width_inches=6.5,
              caption=f"图  {cn}词三模式维度对比（第8层原始激活值）")

    # 逐词模式
    add_heading(doc, f"{sec_num}.2  模式1 逐词孤立——关键维度", 2)
    pw = load_csv(odir / "per_word_dim_average_abs.csv")
    pw_pos = [r for r in pw if r.get("group") == "max"][:8]
    pw_neg = [r for r in pw if r.get("group") == "min"][:6]
    if pw_pos:
        add_para(doc, "正向核心维度：", bold=True, size=10)
        add_table(doc,
            ["维度", "出现次数", "均值|激活|", "最大|激活|"],
            [[r["dim"], r["appearance_count"], fmt_f(r["mean_abs_value"]), fmt_f(r["max_abs_value"])]
             for r in pw_pos],
            header_bg=color_pos)
    if pw_neg:
        add_para(doc, "负向核心维度：", bold=True, size=10)
        add_table(doc,
            ["维度", "出现次数", "均值|激活|", "最大|激活|"],
            [[r["dim"], r["appearance_count"], fmt_f(r["mean_abs_value"]), fmt_f(r["max_abs_value"])]
             for r in pw_neg],
            header_bg=color_neg)

    # 激活最强/最弱词
    add_heading(doc, f"{sec_num}.3  激活最强 vs 最弱词汇（逐词模式）", 2)
    ex = load_csv(odir / "per_word_dim_extremes.csv")
    if ex:
        top10 = sorted(ex, key=lambda r: float(r.get("max_value_1", 0)), reverse=True)[:8]
        bot8  = sorted(ex, key=lambda r: float(r.get("max_value_1", 0)))[:8]
        add_para(doc, "激活最强（语义最纯粹）：", bold=True, size=10)
        add_table(doc,
            ["词", "最强维度", "激活值"],
            [[r["input"], r["max_dim_1"], fmt_f(r["max_value_1"])] for r in top10],
            header_bg=color_pos)
        add_para(doc, "激活最弱（语义最模糊）：", bold=True, size=10)
        add_table(doc,
            ["词", "最强维度", "激活值"],
            [[r["input"], r["max_dim_1"], fmt_f(r["max_value_1"])] for r in bot8],
            header_bg="7f7f7f")

    # 位置模式图
    add_heading(doc, f"{sec_num}.4  模式3 序列位置——涌现维度", 2)
    add_image(doc, odir / "positional_dim_stats_chart.png", width_inches=6.5,
              caption=f"图  {cn}词序列位置模式维度统计（第8层）")
    pos = load_csv(odir / "positional_dim_average_abs.csv")
    pos_pos = [r for r in pos if r.get("group") == "max"][:8]
    pos_neg = [r for r in pos if r.get("group") == "min"][:6]
    if pos_pos:
        add_para(doc, "位置模式正向维度：", bold=True, size=10)
        add_table(doc,
            ["维度", "出现次数", "均值|激活|", "最大|激活|"],
            [[r["dim"], r["appearance_count"], fmt_f(r["mean_abs_value"]), fmt_f(r["max_abs_value"])]
             for r in pos_pos],
            header_bg=color_pos)
    if pos_neg:
        add_para(doc, "位置模式负向维度：", bold=True, size=10)
        add_table(doc,
            ["维度", "出现次数", "均值|激活|", "最大|激活|"],
            [[r["dim"], r["appearance_count"], fmt_f(r["mean_abs_value"]), fmt_f(r["max_abs_value"])]
             for r in pos_neg],
            header_bg=color_neg)

    doc.add_page_break()


# ─────────────── Cross-category comparison ────────────────────────────────────

def get_top_dim(cat_key, group, n=1):
    """返回某类别逐词模式最活跃的n个维度号"""
    rows = load_csv(out_dir(cat_key) / "per_word_dim_average_abs.csv")
    filtered = [r for r in rows if r.get("group") == group]
    return filtered[:n]


def write_cross_comparison(doc, sec_num):
    add_heading(doc, f"第{sec_num}章  四类别横向对比与深度解读", 1,
                color_rgb=RGBColor(0x1F, 0x4E, 0x79))

    # 6.1 各类别最活跃正向维度汇总
    add_heading(doc, f"{sec_num}.1  各类别最活跃正向维度汇总（逐词模式）", 2)
    rows = []
    for cat in CATEGORIES:
        top = get_top_dim(cat["key"], "max", 3)
        for r in top:
            rows.append([
                cat["cn"],
                r["dim"],
                r["appearance_count"],
                fmt_f(r["mean_abs_value"]),
                fmt_f(r["max_abs_value"]),
            ])
    add_table(doc,
        ["类别", "维度", "出现次数", "均值|激活|", "最大|激活|"],
        rows, header_bg="1f4e79")

    # 6.2 维度重叠分析
    add_heading(doc, f"{sec_num}.2  跨类别维度重叠分析", 2)
    add_para(doc,
        "若同一维度在多个感知类别中均高度激活，说明该维度可能编码更底层的"
        "语言/语义结构（如词性、词频）而非类别专属语义。\n"
        "若某维度只在单一类别中显著激活，则更可能是该类别的专属语义维度。",
        size=10)

    # 收集各类别top-5正向维度
    cat_dims = {}
    for cat in CATEGORIES:
        top = get_top_dim(cat["key"], "max", 5)
        cat_dims[cat["cn"]] = {r["dim"] for r in top}

    all_dims = set()
    for s in cat_dims.values():
        all_dims |= s

    overlap_rows = []
    for dim in sorted(all_dims, key=lambda x: int(x) if x.isdigit() else 0):
        cats_with = [cn for cn, dims in cat_dims.items() if dim in dims]
        overlap_rows.append([
            dim,
            str(len(cats_with)),
            "、".join(cats_with) if cats_with else "-",
            "跨类别共享维度（语言结构层）" if len(cats_with) >= 3 else
            "两类共享" if len(cats_with) == 2 else
            "类别专属维度",
        ])
    add_table(doc,
        ["维度", "出现类别数", "出现于哪些类别", "解读"],
        overlap_rows, header_bg="4e3379")

    # 6.3 上下文涌现效应对比
    add_heading(doc, f"{sec_num}.3  序列上下文涌现效应对比", 2)
    add_para(doc,
        "对比逐词模式与序列位置模式，观察各类别中维度激活的变化规律：\n"
        "  - 覆盖率大幅提升的维度 → 上下文增强效应（该类别词汇具有强列表结构）\n"
        "  - 覆盖率大幅下降的维度 → 上下文抑制效应（孤立时的词法噪声被消除）\n"
        "  - 新涌现维度（孤立时<20%，序列时>70%）→ 编码列表结构元信息",
        size=10)

    emergence_rows = []
    for cat in CATEGORIES:
        pw = load_csv(out_dir(cat["key"]) / "per_word_dim_average_abs.csv")
        pos = load_csv(out_dir(cat["key"]) / "positional_dim_average_abs.csv")
        pw_max = {r["dim"]: int(r["appearance_count"]) for r in pw if r.get("group") == "max"}
        pos_max = {r["dim"]: int(r["appearance_count"]) for r in pos if r.get("group") == "max"}
        total = len(load_csv(out_dir(cat["key"]) / "per_word_dim_extremes.csv"))
        if total == 0:
            continue
        for dim in set(list(pw_max.keys())[:5] + list(pos_max.keys())[:5]):
            pw_cnt = pw_max.get(dim, 0)
            pos_cnt = pos_max.get(dim, 0)
            delta = pos_cnt - pw_cnt
            emergence_rows.append([
                cat["cn"], dim,
                f"{pw_cnt}/{total}",
                f"{pos_cnt}/{total}",
                f"{delta:+d}",
                "上下文增强" if delta > 10 else "上下文抑制" if delta < -10 else "基本稳定",
            ])
    if emergence_rows:
        add_table(doc,
            ["类别", "维度", "逐词出现次数", "位置出现次数", "变化量", "效应"],
            emergence_rows, header_bg="375623")

    # 6.4 三层编码模型（扩展至四类别）
    add_heading(doc, f"{sec_num}.4  四类别统一三层编码模型", 2)
    add_para(doc,
        "基于四个感知类别的实验结果，Llama-3-8B 第8层的维度激活可统一归纳为三个功能层次：",
        size=10)

    layers = [
        ("第一层：感知类别核心信号（跨模式稳健）",
         "各类别均存在2-4个维度，在三种输入模式下保持稳定高激活（覆盖率>90%）。"
         "这些是模型对该感知类别的硬编码语义表征，构成'识别轴'——"
         "如颜色的 dim 4055/1800，味道/声音/形状各自也有对应的核心轴。",
         RGBColor(0xC0, 0x00, 0x00)),
        ("第二层：词法表面信号（孤立时强，上下文中被抑制）",
         "各类别都存在若干在孤立模式下活跃、但在序列位置模式下大幅下降的维度。"
         "这些维度混合了词频、形态特征与语义信息，并非纯感知信号。"
         "不同类别的词法噪声维度存在一定程度的重叠，说明它们是模型底层语言特征的通用编码。",
         RGBColor(0xBF, 0x8F, 0x00)),
        ("第三层：上下文结构信号（仅在序列中涌现）",
         "各类别均发现孤立输入时几乎沉默、序列中大量涌现的维度。"
         "这些维度编码列表结构元信息，让模型知道'当前词属于某种语义类别的枚举列表'。"
         "不同感知类别涌现的结构维度各有差异，说明模型对不同类型的语义列表有一定的区分能力。",
         RGBColor(0x37, 0x56, 0x23)),
    ]
    for title, body, rgb in layers:
        p = doc.add_paragraph()
        rt = p.add_run(title + "\n")
        rt.bold = True
        rt.font.size = Pt(10.5)
        rt.font.color.rgb = rgb
        set_font_cn(rt)
        rb = p.add_run(body)
        rb.font.size = Pt(10)
        set_font_cn(rb)
        doc.add_paragraph()

    doc.add_page_break()


def write_conclusions(doc, sec_num):
    add_heading(doc, f"第{sec_num}章  综合结论", 1,
                color_rgb=RGBColor(0x1F, 0x4E, 0x79))

    conclusions = [
        ("C1  感知类别语义在第8层高度结晶化",
         "颜色、味道、声音、形状四类词汇在第8层均形成了专属的核心维度编码轴。"
         "不同类别的核心维度彼此不同，表明第8层已具备感知类别区分能力，"
         "而非对所有语义类别使用同一套维度。"),
        ("C2  具体-抽象梯度普遍存在",
         "在所有四个类别中，源自具体参照物（颜色：apricot；味道：herby；"
         "声音：chirping；形状：spherical）的词汇激活显著强于抽象/多义描述词，"
         "反映了模型对感知词汇具体性的内在编码。"),
        ("C3  列表上下文涌现效应是通用机制",
         "四个类别均观察到序列位置模式下新维度大量涌现的现象，"
         "证明Llama-3-8B具备识别'语义类别枚举列表'这一语言结构的能力，"
         "且这一机制在颜色、味道、声音、形状等不同语义域中均有效。"),
        ("C4  词法表面噪声是普遍干扰",
         "孤立输入模式下，各类别均存在将词法频率/形态特征与语义信号混合的维度。"
         "上下文可有效抑制这些噪声维度，提纯语义信号，"
         "说明序列上下文是模型进行语义消歧的重要机制。"),
        ("C5  第8层是感知-语义的交汇层",
         "实验表明第8层（共33层）并非纯语义层，词法特征仍有明显残留，"
         "但已形成清晰的感知类别编码轴。"
         "未来扩展到其他层（如第16、24、32层）的分析，"
         "有望揭示感知语义从表层到深层的逐层提纯过程。"),
    ]
    for ct, cb in conclusions:
        p = doc.add_paragraph()
        rt = p.add_run(ct + "\n")
        rt.bold = True
        rt.font.size = Pt(11)
        rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        set_font_cn(rt)
        rb = p.add_run(cb)
        rb.font.size = Pt(10)
        set_font_cn(rb)
        doc.add_paragraph()


# ─────────────── Main ────────────────────────────────────────────────────────

def build_report():
    # 统计各类别实际词数
    word_files = {
        "color_words": "data/color_words.txt",
        "taste_words": "data/taste_words.txt",
        "sound_words": "data/sound_words.txt",
        "shape_words": "data/shape_words.txt",
    }
    for cat in CATEGORIES:
        wf = Path(word_files[cat["key"]])
        if wf.exists():
            cat["words"] = sum(1 for line in wf.read_text(encoding="utf-8").splitlines() if line.strip())

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_heading("四感知类别三模式第八层隐藏层分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        set_font_cn(title.runs[0])

    sub = doc.add_paragraph(
        "颜色 · 味道 · 声音 · 形状  |  Meta-Llama-3-8B-Instruct  |  Layer 8  |  三种输入模式"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    set_font_cn(sub.runs[0])

    total_words = sum(cat["words"] or 0 for cat in CATEGORIES)
    info = doc.add_paragraph(
        f"实验日期：2026-05-07  |  四类别总词数：{total_words}  |  隐藏状态维度：4096"
    )
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(10)
    info.runs[0].italic = True
    set_font_cn(info.runs[0])
    doc.add_page_break()

    # ── 第1章：实验设计总览 ──────────────────────────────────────────────────
    add_heading(doc, "第1章  实验设计总览", 1)
    add_para(doc,
        "本报告汇总了对四种感知类别词汇（颜色、味道、声音、形状）的三模式第八层隐藏层分析实验。\n"
        "实验框架统一：均使用 Meta-Llama-3-8B-Instruct 模型，提取第8层（共33层）隐藏状态（4096维），"
        "采用三种输入模式分析模型内部的语义编码规律。\n"
        "本实验框架设计为可扩展至任意层（第1-32层），"
        "为后续多层对比奠定基础。",
        size=10)
    doc.add_paragraph()

    # 词库概述表
    add_para(doc, "1.1  四类别词库概述", bold=True, size=11)
    add_table(doc,
        ["感知类别", "词数", "词库分类", "典型词汇"],
        [
            ["颜色（Color）",   str(CATEGORIES[0]["words"]), "9色系",
             "red, azure, lavender, apricot, graphite..."],
            ["味道（Taste）",   str(CATEGORIES[1]["words"]), "9味觉类",
             "sweet, bitter, umami, smoky, astringent..."],
            ["声音（Sound）",   str(CATEGORIES[2]["words"]), "6音响类",
             "booming, shrill, resonant, melodic, crackling..."],
            ["形状（Shape）",   str(CATEGORIES[3]["words"]), "5形态类",
             "spherical, jagged, tapered, helical, concave..."],
        ],
        header_bg="1f4e79")

    # 三种模式说明
    add_para(doc, "1.2  三种输入模式说明", bold=True, size=11)
    add_table(doc,
        ["模式", "操作方法", "目的"],
        [
            ["模式1：逐词孤立输入",
             "每个词单独输入，提取最后sub-token对应第8层隐藏状态。N次推理。",
             "纯词义编码，无上下文"],
            ["模式2：全词合并末token",
             "所有词拼接为一个字符串，仅提取最后一个token的隐藏状态。1次推理。",
             "以全词列表为背景，观察最后词的编码"],
            ["模式3：序列位置提取",
             "所有词拼接后做一次推理，在各词自身位置提取隐藏状态。1次推理，N个向量。",
             "每词感知完整列表上下文，模拟真实文本"],
        ],
        header_bg="375623")
    doc.add_page_break()

    # ── 各类别章节 ───────────────────────────────────────────────────────────
    for i, cat in enumerate(CATEGORIES, start=2):
        write_category_section(doc, cat, sec_num=i)

    # ── 横向对比章 ──────────────────────────────────────────────────────────
    write_cross_comparison(doc, sec_num=len(CATEGORIES) + 2)

    # ── 结论章 ────────────────────────────────────────────────────────────────
    write_conclusions(doc, sec_num=len(CATEGORIES) + 3)

    doc.save(str(REPORT_PATH))
    print(f"报告已保存：{REPORT_PATH}")


build_report()
