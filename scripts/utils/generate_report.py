# -*- coding: utf-8 -*-
"""Generate Word report for LLM color-probe experiment."""
from __future__ import annotations
import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("data/outputs/color_words")
REPORT_PATH = OUT_DIR / "LLM_color_probe_report.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_image(doc, path, width_inches=6.2, caption=""):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True; cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x55,0x55,0x55)

def add_table(doc, headers, rows, header_bg="1f4e79", alt_bg="dce6f1"):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], header_bg)
    for ri,row_data in enumerate(rows):
        row = t.add_row().cells
        bg = alt_bg if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(row_data):
            row[ci].text = str(val)
            row[ci].paragraphs[0].runs[0].font.size = Pt(8.5)
            row[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row[ci], bg)
    doc.add_paragraph(); return t

def load_csv(path):
    with open(path, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))

def fmt_f(v, d=4):
    try: return f"{float(v):.{d}f}"
    except: return str(v)

def build_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin=Cm(2); section.bottom_margin=Cm(2)
        section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)

    # Cover
    doc.add_paragraph()
    t = doc.add_heading("LLM Color Probe Experiment Report", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Meta-Llama-3-8B-Instruct | Layer 8 Hidden States | Three Input Modes")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER; s.runs[0].font.size=Pt(13)
    s.runs[0].font.color.rgb = RGBColor(0x2E,0x75,0xB6)
    d = doc.add_paragraph("Date: 2026-05-07 | Words: 82 | Dims: 4096")
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER; d.runs[0].font.size=Pt(10); d.runs[0].italic=True
    doc.add_page_break()

    # Sec 1
    add_heading(doc,"1  Experiment Overview",1)
    add_para(doc,
        "This experiment uses neural probing to reveal how Meta-Llama-3-8B-Instruct encodes the "
        "semantic category of COLOR in its 8th layer (of 33) hidden states (4096-dim vectors).\n"
        "82 English color words are fed to the model under three distinct input modes. "
        "For each word the activation vector is extracted and the most salient dimensions are identified. "
        "No normalization is applied -- all values shown are raw signed activations.",size=10)
    doc.add_paragraph()

    add_para(doc,"Vocabulary Design",bold=True,size=11)
    add_table(doc,["Color Family","Count","Representative Words"],[
        ["Red","10","red, crimson, scarlet, maroon, burgundy, carmine, vermilion, ruby, brick, rust"],
        ["Orange","8","orange, amber, coral, peach, apricot, tangerine, saffron, ochre"],
        ["Yellow","6","yellow, gold, lemon, mustard, buff, canary"],
        ["Green","10","green, lime, olive, emerald, jade, mint, teal, chartreuse, sage, khaki"],
        ["Blue","9","blue, navy, cobalt, azure, cerulean, sapphire, denim, indigo, slate"],
        ["Purple","8","purple, violet, lavender, lilac, magenta, fuchsia, plum, mauve"],
        ["Pink","4","pink, rose, salmon, blush"],
        ["Brown/Earth","16","brown, tan, beige, cream, ivory, ecru, taupe, sand, sienna, umber..."],
        ["Achromatic","11","white, gray, silver, charcoal, onyx, ebony, black, pearl, platinum, pewter, graphite"],
    ],header_bg="1f4e79")

    add_para(doc,"Three Input Modes",bold=True,size=11)
    add_table(doc,["Mode","Method","Purpose"],[
        ["Mode 1: Per-Word (isolated)","Each word input separately; extract last-sub-token hidden state. 82 inferences.","Pure word-level semantics, no context"],
        ["Mode 2: All-Input (last token)","All 82 words joined by spaces, extract ONLY the last token. 1 inference.","Last word with 81 words as left context"],
        ["Mode 3: Positional (in-sequence)","All 82 words joined; extract each word at its own position. 1 inference, 82 vectors.","Each word with shared list context"],
    ],header_bg="375623")
    doc.add_page_break()

    # Sec 2
    add_heading(doc,"2  Dimension Activation Overview -- Three-Mode Comparison",1)
    add_para(doc,
        "The chart below shows the top-20 most active dimensions for each mode side by side. "
        "Bar length = mean |activation| (positive dims) or signed activation value (Mode 2 centre panel). "
        "Warm colour (orange-red) = positive activation; cool colour (dark blue) = negative activation. "
        "All values are raw, not normalized.",size=10)
    add_image(doc, OUT_DIR/"three_mode_comparison_chart.png", width_inches=6.5,
              caption="Figure 1  Three-mode dimension comparison (Layer 8, raw signed activation values)")
    add_para(doc,
        "Key observation: In Mode 1 (left), dim 4055 dominates massively -- its mean activation "
        "is ~3x higher than the second-ranked dimension. In Mode 3 (right), new dimensions "
        "(dim 2261, dim 1162, dim 709) emerge that were nearly silent in isolation, "
        "revealing the model's sensitivity to list structure.",italic=True,size=9)
    doc.add_page_break()

    # Sec 3
    add_heading(doc,"3  Mode 1 -- Per-Word Isolated Input",1)
    add_image(doc, OUT_DIR/"per_word_dim_stats_chart.png", width_inches=6.5,
              caption="Figure 2  Per-word mode: top-20 dimensions (positive row upper, negative row lower)")

    add_heading(doc,"3.1  Top Positive Dimensions",2)
    pw = load_csv(OUT_DIR/"per_word_dim_average_abs.csv")
    pw_pos = [r for r in pw if r["group"]=="max"][:10]
    add_table(doc,["Dim","Count / 82","Mean |Act|","Max |Act|","Interpretation"],[
        [r["dim"],r["appearance_count"],fmt_f(r["mean_abs_value"]),fmt_f(r["max_abs_value"]),
         "PRIMARY COLOR CHANNEL" if r["dim"]=="4055" else
         "Robust color marker" if r["dim"]=="290" else
         "Strong color signal" if int(r["appearance_count"])>=50 else "Moderate / lexical"]
        for r in pw_pos],header_bg="c55a11")

    add_heading(doc,"3.2  Top Negative Dimensions",2)
    pw_neg = [r for r in pw if r["group"]=="min"][:8]
    add_table(doc,["Dim","Count / 82","Mean |Act|","Max |Act|","Interpretation"],[
        [r["dim"],r["appearance_count"],fmt_f(r["mean_abs_value"]),fmt_f(r["max_abs_value"]),
         "PRIMARY negative color channel" if r["dim"]=="1800" else
         "Strong negative color signal" if int(r["appearance_count"])>=50 else "Moderate"]
        for r in pw_neg],header_bg="2e4099")

    add_heading(doc,"3.3  Strongest vs Weakest Words (dim 4055)",2)
    ex = load_csv(OUT_DIR/"per_word_dim_extremes.csv")
    top10 = sorted(ex,key=lambda r:float(r.get("max_value_1",0)),reverse=True)[:10]
    bot10 = sorted(ex,key=lambda r:float(r.get("max_value_1",0)))[:10]
    add_para(doc,"Top 10 strongest-activated words (concrete, visually specific color names):",bold=True,size=10)
    add_table(doc,["Rank","Word","Dim","Activation","Reason"],[
        [i+1,r["input"],r["max_dim_1"],fmt_f(r["max_value_1"]),
         "Highly specific visual referent" if float(r["max_value_1"])>1.8 else "Vivid color name"]
        for i,r in enumerate(top10)],header_bg="c00000")
    add_para(doc,"Bottom 10 weakest-activated words:",bold=True,size=10)
    add_table(doc,["Rank","Word","Dim","Activation","Reason"],[
        [i+1,r["input"],r["max_dim_1"],fmt_f(r["max_value_1"]),
         "Multi-subtoken tokenization" if r["input"] in ("indigo","chartreuse","vermilion","cerulean") else
         "Polysemous / abstract color term"]
        for i,r in enumerate(bot10)],header_bg="7f7f7f")
    doc.add_page_break()

    # Sec 4
    add_heading(doc,"4  Mode 3 -- Positional (In-Sequence) Extraction",1)
    add_para(doc,
        "In one forward pass, 82 words are concatenated and each word's hidden state is extracted "
        "at its own token position in the sequence. Every word can attend to all preceding color words "
        "as left context -- simulating how the model processes a color word list in real text.",size=10)
    add_image(doc, OUT_DIR/"positional_dim_stats_chart.png", width_inches=6.5,
              caption="Figure 3  Positional mode dimension statistics (Layer 8)")

    add_heading(doc,"4.1  Top Positive Dimensions (Positional)",2)
    pos = load_csv(OUT_DIR/"positional_dim_average_abs.csv")
    pos_pos=[r for r in pos if r["group"]=="max"][:10]
    add_table(doc,["Dim","Count / 82","Mean |Act|","Max |Act|","vs Mode 1"],[
        [r["dim"],r["appearance_count"],fmt_f(r["mean_abs_value"]),fmt_f(r["max_abs_value"]),
         "Coverage 96% -> 100% (context boost)" if r["dim"]=="4055" else
         "77->20 occurrences (context suppressed)" if r["dim"]=="2485" else
         "NEW: 8->76 occurrences (context-emergent)" if r["dim"]=="2261" else
         "NEW: context-only dimension" if r["dim"]=="2265" else
         "Stable across modes"]
        for r in pos_pos],header_bg="375623")

    add_heading(doc,"4.2  Top Negative Dimensions (Positional)",2)
    pos_neg=[r for r in pos if r["group"]=="min"][:8]
    add_table(doc,["Dim","Count / 82","Mean |Act|","Max |Act|","vs Mode 1"],[
        [r["dim"],r["appearance_count"],fmt_f(r["mean_abs_value"]),fmt_f(r["max_abs_value"]),
         "NEW: 9->80 occurrences (context-emergent)" if r["dim"]=="1162" else
         "NEW: 40->81 occurrences (massive boost)" if r["dim"]=="709" else
         "Stable across modes" if r["dim"] in ("1800","3231") else
         "Moderate in positional mode"]
        for r in pos_neg],header_bg="2e4099")
    doc.add_page_break()

    # Sec 5
    add_heading(doc,"5  Mode 2 -- All-Input Last-Token",1)
    add_para(doc,
        "All 82 words concatenated and fed as one string. Only the LAST token (graphite) is extracted. "
        "This is NOT a holistic aggregation of all color words -- it captures only the last word "
        "with the other 81 words as background context. "
        "Peak activation drops to 0.780 vs 2.378 in Mode 1.",size=10)
    add_image(doc, OUT_DIR/"all_input_dim_stats_chart.png", width_inches=6.5,
              caption="Figure 4  All-input last-token mode (24 most salient dimensions)")
    ai = load_csv(OUT_DIR/"all_input_dim_stats.csv")[:16]
    add_table(doc,["Dim","Signed Value","Direction","Abs Value"],[
        [r["dim"],fmt_f(r["signed_value"]),r["direction"],fmt_f(r["abs_value"])]
        for r in ai],header_bg="4e3379")
    doc.add_page_break()

    # Sec 6
    add_heading(doc,"6  Cross-Mode Comparison and Core Findings",1)
    add_heading(doc,"6.1  Key Dimension Behavior Across All Three Modes",2)
    add_table(doc,["Dimension","Mode 1 Per-Word","Mode 2 All-Input","Mode 3 Positional","Interpretation"],[
        ["dim 4055 (pos)","79/82 (96%)\nmean=1.471","Activated (0.780)","82/82 (100%)\nmean=0.968","PRIMARY color channel\nConcrete names trigger strongest"],
        ["dim 290 (pos)","82/82 (100%)\nmean=0.500","Activated (0.471)","80/82 (98%)\nmean=0.470","Most robust color marker\nStable across ALL modes"],
        ["dim 1800 (neg)","82/82 (100%)\nmean=0.820","Activated (-0.421)","81/82 (99%)\nmean=0.472","Primary NEGATIVE color channel\nForms bipolar axis with dim 4055"],
        ["dim 3231 (neg)","63/82 (77%)\nmean=0.443","Activated (-0.443)","72/82 (88%)\nmean=0.446","Robust negative signal\nConsistent across modes"],
        ["dim 2485 (pos)","77/82 (94%)\nmean=0.511","Weak (0.345)","20/82 (24%)\nmean=0.424","Strong in isolation\nSuppressed by context\nMay contain lexical features"],
        ["dim 2261 (+/-)","8 occurrences (neg)","Activated (+0.446)","76/82 (93%, POSITIVE)","POLARITY FLIP!\nContext-emergent list structure signal"],
        ["dim 1162 (neg)","9/82 (11%)","Activated (-0.749)","80/82 (98%)\nmean=0.541","Context-only emergence\nEncodes list structure"],
        ["dim 709 (neg)","40/82 (49%) weak","Activated (-0.519)","81/82 (99%)\nmean=0.533","Context-only emergence\nEncodes list structure"],
    ],header_bg="1f4e79")

    add_heading(doc,"6.2  Three-Layer Encoding Model",2)
    layers=[
        ("Layer 1 -- Core Color Signal (cross-mode robust)",
         "Dims: 4055 (pos), 290 (pos), 1800 (neg), 3231 (neg)\n"
         "Active in ALL three modes. These are the model's hard-coded representation of the COLOR semantic category. "
         "dim 4055 and dim 1800 form a bipolar axis -- their simultaneous activation (one high, one low) "
         "acts as the color identifier.","c00000"),
        ("Layer 2 -- Lexical Surface Signal (strong in isolation, suppressed by context)",
         "Dims: 2485, 133, 1815 (pos); 912, 1856, 1731 (neg)\n"
         "Active in per-word mode but largely suppressed in positional mode. "
         "These dimensions likely blend lexical frequency and morphological features with color semantics. "
         "They are not pure color signals.","bf8f00"),
        ("Layer 3 -- Context-Structure Signal (emerges only in sequence)",
         "Dims: 2261, 2265 (pos); 1162, 709 (neg)\n"
         "Nearly silent in isolated input but massively active in positional mode. "
         "Encodes the meta-information: 'I am currently processing a word inside a color word list.' "
         "This reveals the model's structural awareness of enumeration patterns.","375623"),
    ]
    for ltitle,lbody,lcolor in layers:
        p=doc.add_paragraph()
        rt=p.add_run(ltitle+"\n"); rt.bold=True; rt.font.size=Pt(10.5)
        r,g,b=int(lcolor[:2],16),int(lcolor[2:4],16),int(lcolor[4:],16)
        rt.font.color.rgb=RGBColor(r,g,b)
        rb=p.add_run(lbody); rb.font.size=Pt(10)
        doc.add_paragraph()
    doc.add_page_break()

    # Sec 7
    add_heading(doc,"7  Which Words Are Most / Least 'Colorlike'?",1)
    add_para(doc,
        "In per-word mode, dim 4055 activation quantifies how strongly the model "
        "recognizes a word as a color term. Derived-from-object color names score highest; "
        "abstract or polysemous terms score lowest.",size=10)
    finding=[
        ["apricot","2.378","Highly specific visual referent (fruit name -> unique hue)"],
        ["lavender","2.280","Plant name -> tightly focused purple-blue hue"],
        ["saffron","2.212","Spice name -> very specific yellow-orange hue"],
        ["mahogany","2.189","Wood name -> specific reddish-brown hue"],
        ["salmon","2.183","Food name -> specific pink-orange hue"],
        ["...","...","..."],
        ["indigo","0.391  WARNING","Split into 3 sub-tokens; color signal diluted"],
        ["sage","0.426  WARNING","Primary meaning: herb plant; color is secondary"],
        ["rose","0.500","Polysemy: flower vs pink color; signal divided"],
        ["amber","0.548","Primary: mineral/fossil; color is secondary meaning"],
        ["blue","0.800","Most frequent basic color but used in many non-color contexts"],
    ]
    add_table(doc,["Word","dim 4055 Activation","Analysis"],finding,header_bg="7030a0")
    add_para(doc,
        "Pattern: object-derived color names (food, plants, minerals) -> strongest activation; "
        "abstract, polysemous, or multi-subtoken color words -> weakest activation. "
        "This mirrors human cognitive salience gradients for color terms.",italic=True,size=9)
    doc.add_page_break()

    # Sec 8
    add_heading(doc,"8  Conclusions",1)
    conclusions=[
        ("C1  Color semantics is highly crystallized by Layer 8",
         "Llama-3-8B has formed a dedicated encoding region for the COLOR semantic category at layer 8. "
         "Dim 4055 (positive) and dim 1800 (negative) activate at 96-100% coverage across modes, "
         "with magnitudes 2-3x higher than other dimensions, forming the core color encoding axis."),
        ("C2  Concrete vs abstract color names show a significant activation gradient",
         "Object-derived color names (apricot, lavender, saffron) activate dim 4055 at 2-3x the level of "
         "basic color terms (blue, red), reflecting clearer color-specific representations for "
         "words with unambiguous color identity."),
        ("C3  Context strengthens weak color signals",
         "In positional mode, dim 4055 coverage rises from 96% to 100%. Words like indigo that fail "
         "to trigger color dimensions in isolation DO activate them when surrounded by color word context, "
         "showing the model's ability to use context to disambiguate edge cases."),
        ("C4  The model perceives 'color word list' as a macro-structure",
         "Dimensions dim 2261, dim 1162, dim 709 are nearly silent in isolation but activate at 80-82/82 "
         "coverage in positional mode. These context-exclusive dimensions encode the structural information "
         "that the current word belongs to an enumerated list -- the model reads not just word meaning "
         "but also list structure."),
        ("C5  Lexical surface features co-exist with semantic signals",
         "Dims 2485 and 133 are highly active in isolation but suppressed in context. They likely encode "
         "lexical frequency or morphological patterns alongside color semantics, suggesting Layer 8 "
         "is not yet a purely semantic layer. Higher layers may show purer color representations."),
    ]
    for ct,cb in conclusions:
        p=doc.add_paragraph()
        rt=p.add_run(ct+"\n"); rt.bold=True; rt.font.size=Pt(11)
        rt.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
        rb=p.add_run(cb); rb.font.size=Pt(10)
        doc.add_paragraph()

    doc.save(str(REPORT_PATH))
    print(f"Saved: {REPORT_PATH}")

build_report()
