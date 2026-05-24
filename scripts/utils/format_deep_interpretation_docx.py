from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "outputs"
SOURCE = OUT / "深度解读.docx"
TARGET = OUT / "深度解读_排版优化版.docx"

ACCENT = "1F4E79"
ACCENT_2 = "C55A11"
SOFT_BLUE = "EAF2F8"
SOFT_ORANGE = "FCE4D6"
SOFT_GREEN = "E2F0D9"
GRID = "D9E2F3"
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x55, 0x65, 0x73)


def set_cell_bg(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D0D7DE", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def style_run(run, size=10.5, bold=False, color=TEXT) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_para(doc: Document, text: str, size=10.5, bold=False, color=TEXT, before=0, after=6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 4)
    run = p.add_run(text)
    style_run(run, size=16 if level == 1 else 12.5, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))


def add_callout(doc: Document, title: str, body: str, fill: str = SOFT_BLUE, accent: str = ACCENT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    cell = table.cell(0, 0)
    set_cell_bg(cell, fill)
    set_cell_border(cell, accent, "10")
    set_cell_margins(cell, top=150, start=170, bottom=150, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    style_run(r, size=11, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    style_run(r2, size=10.2, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = hdr[idx]
        set_cell_bg(cell, ACCENT)
        set_cell_border(cell, "FFFFFF")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        style_run(r, size=9.2, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, text in enumerate(row):
            cell = cells[col_idx]
            set_cell_bg(cell, "FFFFFF" if row_idx % 2 == 0 else "F8FAFC")
            set_cell_border(cell, "D7DEE8")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            style_run(r, size=8.8, color=TEXT)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("四感知类别三模式第八层深度解读  |  Meta-Llama-3-8B-Instruct")
    style_run(run, size=8.5, color=MUTED)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(56)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("四感知类别三模式第八层分析\n深度解读")
    style_run(r, size=24, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    r = subtitle.add_run("颜色 / 味道 / 声音 / 形状  |  Layer 8 Hidden States")
    style_run(r, size=12.5, color=RGBColor(0xC5, 0x5A, 0x11))

    add_callout(
        doc,
        "阅读提示",
        "本版重点优化阅读层级：先给出核心判断，再展示关键数据，最后给出方法论反思。"
        "维度不被直接等同于人类概念，而是解释为高维空间中的功能性方向。",
        fill=SOFT_BLUE,
    )

    add_table(
        doc,
        ["分析对象", "词数", "输入模式", "探针层", "核心读法"],
        [["四类感知词", "293", "逐词 / 全量末 token / 序列位置", "第 8 层", "维度方向 + 覆盖率 + 平均激活"]],
        widths=[2.4, 1.4, 4.5, 1.6, 5.0],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1])

    add_heading(doc, "1. 核心发现：dim 4055 是跨感知类别的语义识别轴", 1)
    add_callout(
        doc,
        "一句话结论",
        "dim 4055 不是颜色专属维度，而更像第 8 层对“具体感知属性描述词”的通用强激活通道。",
        fill=SOFT_ORANGE,
        accent=ACCENT_2,
    )
    add_table(
        doc,
        ["类别", "词数", "逐词均值", "逐词覆盖", "位置模式均值", "位置模式覆盖"],
        [
            ["颜色", "82", "1.471", "79/82 (96%)", "0.968", "82/82 (100%)"],
            ["味道", "73", "1.453", "72/73 (99%)", "1.074", "73/73 (100%)"],
            ["声音", "69", "1.358", "69/69 (100%)", "1.098", "69/69 (100%)"],
            ["形状", "69", "1.369", "69/69 (100%)", "0.936", "69/69 (100%)"],
        ],
        widths=[1.7, 1.2, 2.0, 2.5, 2.4, 2.8],
    )
    add_para(
        doc,
        "解读：dim 4055 对四个完全不同的感知类别均以高覆盖率、高幅度激活。"
        "因此它不应被命名为“颜色维度”，而应被理解为模型对“可感知属性词汇”这一宏观语义类别的通用探测器。"
        "凡是描述看、尝、听、触/形等感觉经验的词汇，都容易在该方向上留下强信号。",
    )

    add_heading(doc, "2. dim 1800 与 dim 4055 构成双极编码轴", 1)
    add_table(
        doc,
        ["类别", "覆盖率", "平均绝对值", "解读重点"],
        [
            ["颜色", "82/82 (100%)", "0.820", "颜色词语义最稳定，负向轴印记最强"],
            ["味道", "73/73 (100%)", "0.674", "味觉词也形成稳定负向对照信号"],
            ["声音", "69/69 (100%)", "0.587", "声音词覆盖稳定，但幅度略弱"],
            ["形状", "69/69 (100%)", "0.674", "形状词与味觉词幅度接近"],
        ],
        widths=[1.7, 2.4, 2.2, 7.4],
    )
    add_para(
        doc,
        "dim 4055（正向）与 dim 1800（负向）同时稳定出现，可以被看作感知语义空间中的双极轴。"
        "这并不意味着 dim 1800 表示“反颜色”或负面意义，而是说明模型通过一组正负方向共同定位词的感知属性强弱。"
    )

    add_heading(doc, "3. 三层编码结构：语义、词法、上下文", 1)
    add_table(
        doc,
        ["层次", "代表维度", "主要现象", "功能解释"],
        [
            ["核心感知轴", "4055 / 290 / 1800", "跨类别、跨模式稳定", "感知属性词的中层语义识别"],
            ["词法表面噪声", "2485 / 1815 / 912 / 1856", "孤立输入强，序列位置中被抑制", "词频、词形、英文 token 规范性等混合信号"],
            ["列表结构信号", "2261 / 2265 / 1162", "孤立时几乎沉默，列表中涌现", "模型识别到“同类词枚举列表”的语用结构"],
        ],
        widths=[2.4, 3.2, 4.2, 4.8],
    )
    add_para(doc, "词法表面噪声的典型抑制如下：", bold=True)
    add_table(
        doc,
        ["维度", "颜色 pw→pos", "味道 pw→pos", "声音 pw→pos", "形状 pw→pos"],
        [
            ["dim 2485", "77→20/82", "72→11/73", "69→23/69", "69→41/69"],
            ["dim 1815", "62→9/82", "56→15/73", "53→14/69", "55→9/69"],
        ],
        widths=[2.2, 3.0, 3.0, 3.0, 3.0],
    )
    add_para(doc, "上下文列表结构信号的涌现如下：", bold=True)
    add_table(
        doc,
        ["类别", "dim 2261 pw→pos", "dim 2265 pw→pos", "dim 1162 pw→pos"],
        [
            ["颜色", "0→76/82", "0→74/82", "9→80/82"],
            ["味道", "0→63/73", "0→69/73", "5→64/73"],
            ["声音", "0→50/69", "0→62/69", "13→66/69"],
            ["形状", "0→61/69", "0→51/69", "≈0→68/69"],
        ],
        widths=[1.8, 3.5, 3.5, 3.5],
    )

    add_heading(doc, "4. 激活强弱梯度：词义越专一，激活越强", 1)
    add_table(
        doc,
        ["类别", "强激活词示例", "弱激活词示例", "语言学解释"],
        [
            ["颜色", "apricot / lavender / saffron", "indigo / rose / amber", "具体色名强，多义或多子词弱"],
            ["味道", "rancid / spicy / fermented", "fresh / hot / bold / clean", "极端味觉体验强，跨域形容词弱"],
            ["声音", "guttural / raucous / shrill", "echoing / clear / flat", "特化音质词强，通用描述词弱"],
            ["形状", "concave / prismatic / octagonal", "oval / short / broad", "几何术语强，日常宽泛词弱"],
        ],
        widths=[1.6, 4.0, 3.6, 5.0],
    )
    add_callout(
        doc,
        "解释框：原型效应",
        "激活强度与词义绑定专一程度高度相关。术语性强、跨域使用少、语义聚焦的词激活最强；日常高频、多域通用、语义分散的词激活最弱。",
        fill=SOFT_GREEN,
        accent="548235",
    )

    add_heading(doc, "5. 四类别差异：同一机制，不同语义纹理", 1)
    for label, body in [
        ("颜色", "核心轴幅度最高，说明颜色词在视觉语言中语义最集中，词义几乎完全由感知特征定义。"),
        ("声音", "dim 4055 在逐词模式达到 100% 覆盖，即使 sharp、flat 等跨域词也被声音语义域捕获。"),
        ("味道", "位置模式中结构性编码更强，说明味觉词列表对上下文结构信号较敏感。"),
        ("形状", "形状词的词法维度抑制较弱，可能因为形状词更概念化、几何化，词形和语义绑定更紧。"),
    ]:
        add_para(doc, f"{label}：{body}", before=2, after=4)

    add_heading(doc, "6. 方法论反思：三模式中最该信任哪一个？", 1)
    add_callout(
        doc,
        "谨慎解释",
        "模式 2（全词末 token）不是全部词的平均表征，而是最后一个词在长上下文中的隐藏状态。因此它只能作为上下文压缩后的末位读数，不能直接代表整个类别。",
        fill=SOFT_ORANGE,
        accent=ACCENT_2,
    )
    add_para(
        doc,
        "模式 1 与模式 3 的对比最有价值：模式 1 显示词本身的孤立语义，模式 3 显示模型在知道自己处理同类词列表时如何改变隐藏状态。"
        "未来可扩展到第 1-4 层与第 24-32 层，验证感知语义是否从词法混合逐步走向类别专属化。",
    )

    add_heading(doc, "7. 总结", 1)
    add_table(
        doc,
        ["结论", "含义"],
        [
            ["第 8 层不是纯语义层", "仍混合词法、token 形态和上下文结构信号"],
            ["感知语义已明显结晶", "4055/1800/290 等维度形成稳定感知轴"],
            ["上下文会重塑维度分布", "2261/2265/1162 等维度在列表模式中涌现"],
            ["不要把单维度等同人类概念", "更可靠的说法是“该维度近似关联某类功能方向”"],
        ],
        widths=[4.0, 10.0],
    )

    doc.save(TARGET)


if __name__ == "__main__":
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    build()
    print(TARGET)
